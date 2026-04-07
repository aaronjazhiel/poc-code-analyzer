import json
import boto3
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

S3_BUCKET = "poc-code-analyzer-2026"
S3_PREFIX = "documentos-generados/asis/"

def lambda_handler(event, context):
    print("AS-IS Lambda iniciada")
    try:
        agent        = event.get("agent", {})
        action_group = event.get("actionGroup", "")
        function     = event.get("function", "")
        parameters   = event.get("parameters", [])
        param_dict   = {p["name"]: p["value"] for p in parameters}

        if function != "generar-documento-asis":
            return _error_response(event, action_group, function, f"Función no reconocida: {function}")

        stack_actual        = param_dict.get("stack_actual",        "No detectado")
        patron_arquitectura = param_dict.get("patron_arquitectura",  "No detectado")
        componentes         = param_dict.get("componentes",          "No detectado")
        dependencias        = param_dict.get("dependencias",         "No detectado")
        riesgos             = param_dict.get("riesgos",              "No detectado")

        doc = Document()

        # Título
        titulo = doc.add_heading("ANÁLISIS AS-IS — ESTADO ACTUAL DEL SISTEMA", 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in titulo.runs:
            run.font.color.rgb = RGBColor(0x0A, 0x2F, 0x6B)

        # Fecha
        fecha = doc.add_paragraph()
        fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fecha.add_run(f"Generado: {datetime.utcnow().strftime('%d/%m/%Y %H:%M:%S')} UTC")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        doc.add_paragraph()

        secciones = [
            ("1. Stack Tecnológico Actual", stack_actual),
            ("2. Patrón Arquitectónico",    patron_arquitectura),
            ("3. Componentes Principales",  componentes),
            ("4. Dependencias Externas",    dependencias),
            ("5. Observaciones y Riesgos",  riesgos),
        ]

        for titulo_sec, contenido in secciones:
            h = doc.add_heading(titulo_sec, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
            p = doc.add_paragraph(contenido)
            p.runs[0].font.size = Pt(11)
            doc.add_paragraph()

        # Footer
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Documento generado automáticamente por el Agente AS-IS | PoC — Plataforma de Análisis Agéntico de Código")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # Guardar en memoria y subir a S3
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        timestamp = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
        s3_key = f"{S3_PREFIX}asis-{timestamp}.docx"

        s3_client = boto3.client("s3")
        s3_client.put_object(
            Bucket=S3_BUCKET,
            Key=s3_key,
            Body=buffer.getvalue(),
            ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        s3_url = f"s3://{S3_BUCKET}/{s3_key}"
        print(f"Documento AS-IS guardado en: {s3_url}")

        resultado = {
            "status": "ok",
            "documento": "AS-IS generado correctamente en formato Word (.docx)",
            "ubicacion": s3_url,
            "resumen": {
                "stack_actual": stack_actual,
                "patron_arquitectura": patron_arquitectura,
                "componentes": componentes,
                "dependencias": dependencias,
                "riesgos": riesgos
            }
        }

        action_response = {
            "agent": agent, "actionGroup": action_group, "function": function,
            "functionResponse": {"responseBody": {"TEXT": {"body": json.dumps(resultado, ensure_ascii=False)}}}
        }
        return {"response": action_response, "messageVersion": event.get("messageVersion", "1.0")}

    except Exception as e:
        print(f"Error: {e}")
        return {"error": str(e)}

def _error_response(event, action_group, function, message):
    return {
        "response": {
            "actionGroup": action_group, "function": function,
            "functionResponse": {"responseBody": {"TEXT": {"body": json.dumps({"error": message})}}}
        },
        "messageVersion": event.get("messageVersion", "1.0")
    }
